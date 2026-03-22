import time
import random
import tempfile
import os
import glob
from functools import wraps
import streamlit as st
import google.generativeai as genai
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd
import PyPDF2
from docx import Document
import io
import numpy as np


# Streamlit 웹 UI 구성

st.set_page_config(page_title="RFP 초안작성 AI 시스템", page_icon="📝", layout="centered")

hide_st_style = """
    <style>
    #MainMenu {visibility: hidden;} /* 구형 MainMenu 대응 */
    header {visibility: hidden;} /* 우측 상단 점 세개 메뉴 숨기기 */
    footer {visibility: hidden;} /* 하단 기본 바 숨기기 */
    div[data-testid="stStatusWidget"] { visibility: hidden; } /* 우측 하단 내 깃허브 배지 숨기기 */
    

    .block-container {padding-top: 1rem; padding-bottom: 0rem;}
    </style>
"""
st.markdown(hide_st_style, unsafe_allow_html=True)

st.title("📝 제안요청서 초안작성 AI 시스템(국립부경대 김진명 作)")
st.markdown("참고 문서를 첨부하고 지시사항을 구체적으로 입력하면 AI가 맞춤형 초안을 작성해 줍니다.")
st.markdown("집중적으로 작성할 파트를 선택하세요(예시 : 요약본, 요구사항 상세)")
st.markdown("※제안요청서 전체를 작성하려고 할 시 에러 발생 가능성 있어 챕터별 초안작성만 제공됩니다")
st.markdown("※제안요청서 작성완료까지는 약 3~10분 내로 소요됩니다!")


# API 키 설정 (로컬 테스트용)
#try:
    # 1. 클라우드 배포 상태일 때
API_KEY = st.secrets["GEMINI_API_KEY"]
#except:
    # 2. 내 PC에서 로컬 테스트할 때
    #API_KEY = "AIzaSyCvrROFm9b_xPAdD6syfPB1dJRlW95w8HA"

# 테스트모드
#API_KEY = "AIzaSyCvrROFm9b_xPAdD6syfPB1dJRlW95w8HA"

# 1. 과거 HUG 제안요청서 PDF 데이터 로드 (캐싱 적용)


@st.cache_data
def load_reference_rfps(folder_path="reference_rfps"):
    pdf_texts = {}
    if not os.path.exists(folder_path):
        return pdf_texts
    
    # 폴더 안의 모든 PDF 파일 경로를 가져옵니다.
    filepaths = glob.glob(os.path.join(folder_path, "*.pdf"))
    for filepath in filepaths:
        filename = os.path.basename(filepath)
        text = ""
        try:
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                # 앞부분(주로 개요 및 요구사항) 위주로 빠르게 읽기 위해 최대 30페이지만 읽음
                for page in reader.pages[:30]: 
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            pdf_texts[filename] = text
        except Exception as e:
            pass # 암호가 걸려있거나 깨진 PDF는 건너뜀
    return pdf_texts


# 2. 문서 유사도 분석 함수 (과거 사업 Top 5 추출)

@st.cache_data(show_spinner=False)
def get_document_embeddings_v2(_corpus_dict, api_key):
    genai.configure(api_key=api_key)
    embeddings = {}
    
    total_docs = len(_corpus_dict)
    if total_docs == 0:
        return embeddings

    # 진행 상태 바
    progress_text = f"과거 사업 데이터 학습 중... (총 {total_docs}개)"
    my_bar = st.progress(0, text=progress_text)
    
    for i, (filename, text) in enumerate(_corpus_dict.items()):
        chunk = text[:3000] 
        if chunk.strip():
            try:
                res = genai.embed_content(
                    model="models/gemini-embedding-001",
                    content=chunk,
                    task_type="retrieval_document"
                )
                embeddings[filename] = res['embedding']
            except Exception as e:
                # 구글 서버에서 API 거절 사유
                st.error(f"[{filename}] 구글 API 거절 사유: {e}") 
                
        time.sleep(4.2) # API 제한 방지
        my_bar.progress((i + 1) / total_docs, text=f"{progress_text} - ({i+1}/{total_docs}) 완료")
        
    my_bar.empty()
    return embeddings

def get_top_5_similar_rfps_rag(query_text, corpus_dict, api_key):
    if not corpus_dict or not query_text:
        return None
    
    genai.configure(api_key=api_key)
    # v2 함수를 호출하도록 변경 (기존 꼬인 캐시 무시)
    doc_embeddings_dict = get_document_embeddings_v2(corpus_dict, api_key)
    
    if not doc_embeddings_dict:
        st.error("문서 임베딩 데이터가 없습니다. API 제한이 걸렸을 수 있습니다.")
        return None

    try:
        query_chunk = query_text[:3000]
        query_res = genai.embed_content(
            model="models/gemini-embedding-001",
            content=query_chunk,
            task_type="retrieval_query"
        )
        query_vector = np.array(query_res['embedding'])
        
        results = []
        for filename, d_vec in doc_embeddings_dict.items():
            d_vec_np = np.array(d_vec)
            cos_sim = np.dot(query_vector, d_vec_np) / (np.linalg.norm(query_vector) * np.linalg.norm(d_vec_np))
            
            clean_name = filename.replace(".pdf", "")
            results.append({
                "유사 과거 사업명": clean_name,
                "원본파일명": filename,
                "유사도(%)": round(float(cos_sim) * 100, 2)
            })
            
        df = pd.DataFrame(results).sort_values(by="유사도(%)", ascending=False).head(5)
        df.index = range(1, len(df) + 1)
        return df
    except Exception as e:
        st.error(f"유사도 계산 중 오류 발생: {e}")
        return None


# 3. 지수 백오프 및 API 호출 함수

def retry_with_exponential_backoff(max_retries=5, base_delay=2.0, max_delay=60.0):
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            for attempt in range(max_retries):
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    error_msg = str(e).lower()
                    if "429" in error_msg or "quota" in error_msg or "resourceexhausted" in error_msg:
                        if attempt == max_retries - 1:
                            st.error(f"최대 재시도 횟수 초과.")
                            raise e
                        delay = min(max_delay, base_delay * (2 ** attempt))
                        jitter = random.uniform(0, 0.1 * delay)
                        sleep_time = delay + jitter
                        st.warning(f"API 대기 중... {sleep_time:.1f}초 후 재시도합니다.")
                        time.sleep(sleep_time)
                    else:
                        raise e
        return wrapper
    return decorator

@retry_with_exponential_backoff()
def generate_draft(api_key, prompt, uploaded_files=None):
    genai.configure(api_key=api_key)
    generation_config = genai.types.GenerationConfig(max_output_tokens=8192, temperature=0.2)
    model = genai.GenerativeModel('gemini-2.5-pro', generation_config=generation_config)
    
    contents = [prompt]
    temp_file_paths = []
    
    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_extension = os.path.splitext(uploaded_file.name)[1].lower()
            if file_extension in ['.txt', '.csv', '.md']:
                try:
                    text_data = uploaded_file.getvalue().decode('utf-8')
                except:
                    text_data = uploaded_file.getvalue().decode('cp949', errors='replace')
                contents.append(f"\n\n--- [참고자료: {uploaded_file.name}] ---\n{text_data}\n-------------------\n")
            elif file_extension == '.pdf':
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    temp_file_paths.append(tmp_file.name)
        try:
            gemini_files = []
            for path in temp_file_paths:
                uploaded_gemini_file = genai.upload_file(path)
                gemini_files.append(uploaded_gemini_file)
            contents.extend(gemini_files)
            response = model.generate_content(contents)
            return response.text
        finally:
            for path in temp_file_paths:
                if os.path.exists(path):
                    os.remove(path)
    else:
        response = model.generate_content(contents)
        return response.text



# 백그라운드에서 과거 RFP 파일들을 미리 로딩
reference_rfps_dict = load_reference_rfps()

uploaded_files = st.file_uploader(
    "기존 제안요청서나 요구사항 정의서를 업로드하세요 (초안 생성 참고용)", 
    type=['pdf', 'txt', 'csv'], 
    accept_multiple_files=True
)

section_choice = st.radio(
    "집중적으로 작성할 파트를 선택하세요:",
    ["전체 요약본 (짧게)", "1~2장. 사업 개요 및 현황", "3장. 요구사항 상세", "4장. 제안 안내 및 평가 기준"],
    horizontal=True
)

user_input = st.text_area("이번 사업의 핵심 요구사항이나 특별히 강조할 내용을 입력하세요. 예산과 과업기간은 꼭 입력해주시기바랍니다", height=100)

if st.button(f"초안 생성 및 과거 유사 사업 탐색", type="primary"):
    if not API_KEY:
        st.error("API 키를 코드에 입력해 주세요.")
    elif not user_input:
        st.warning("작성할 내용을 입력해 주세요.")
    else:
        with st.spinner("AI가 초안을 작성하고, 사내 과거 사업 DB와 유사도를 비교 중입니다..."):
            try:
                SYSTEM_PROMPT = f"""
                당신은 주택도시보증공사(HUG)의 관리직 직원입니다.
                주어진 과업을 추진하기 위해 제안요청서 초안을 작성해야 합니다.

                [작성 규칙]
                1. 반드시 명확하고 간결한 '개조식(~함, ~임)'으로 작성할 것.
                2. 예산, 일정, 요구사항 등 수치나 명확한 팩트가 있다면 표 형식으로 정리할 것.
                3. 규정에 부합하는 용어와 공공기관 행정 용어를 적절히 사용할 것.
                4. 공공기관 공통 필수 요건 : 사업 분야를 막론하고 공공기관 용역에 공통으로 적용되는 다음 항목들을 문맥에 맞게 적절히 포함할것.
                   - 과업 수행에 따른 산출물 제출 및 검수 기준
                   - 과업 수행 중 발생하는 산출물에 대한 지적재산권(저작권) 귀속 및 보안/비밀유지 의무
                   - 과업 지연 또는 불량에 따른 책임 및 손해배상(페널티) 조건
                5. 할루시네이션(환각) 금지: 사용자가 제공하지 않은 구체적인 예산 금액이나 가상의 법령을 임의로 지어내지 말것.
                6. 내용 축약 및 얼버무리기 절대 금지: 표의 '요구사항 상세설명'이나 '과업 내용'을 작성할 때, 절대 '세부내용'과 같은 단어 하나로 대충 얼버무리거나 축약하지 말것. 정보화 사업(감리, 개인정보영향평가, 컨설팅 용역 등 포함)의 경우 실제 개발 및 사업 수행에 필요한 구체적인 스펙과 동작 방식을 최소 3~4문장 이상의 긴 호흡으로 상세하고 빼곡하게 모두 서술.
                7. 너무 무리하게 분량을 늘리다가 문서가 중간에 끊기지 않도록, 도출하는 항목의 수(약 10~15개 내외)를 조절하여 반드시 기승전결이 있는 완벽하게 끝맺음 된 형태(문장의 마침표 및 표의 닫힘 등)로 출력을 완료할것.
                
                [HUG 표준 목차 (참고용 배경지식)]
                1. 사업 개요 / 2. 공사업무 현황 / 3. 사업 추진방안 / 4. 요구사항 상세 / 5. 제안서 작성요령 / 6. 제안 안내사항

                [HUG 특화 작성 지침]
                1. HUG 비전 및 미션 연계: 본 과업이 단순한 용역을 넘어, HUG의 핵심 미션인 '서민 주거안정', '전세사기 피해 예방', '주택도시기금의 효율적 운용 및 관리' 등에 어떻게 기여할 수 있는지 과업 목적과 기대효과에 자연스럽게 녹여낼것.
                2. 최고 수준의 정보보안 및 데이터 보호: 금융, 보증, 부동산 등 국민의 매우 민감한 재산 및 개인정보를 다루는 공사의 특성상, 용역 수행 과정(행사, 채용, 연구, IT 등 분야 불문)에서 발생할 수 있는 '개인정보 유출 방지 대책' 및 '강력한 보안 서약 요건'을 명시할것.
                3. 대국민 및 유관기관 이해관계자 고려: HUG의 주요 고객인 '일반 국민(임차인/수분양자)', '건설/주택사업자', '금융기관', '국토교통부' 등 복잡한 이해관계자를 고려하여, 용역 결과물이 이들에게 미칠 영향과 소통 계획을 제안서에 포함하도록 요구할것.
                4. 과년도 용역사업과 비슷한 사업을 추진한다면 과년도 용역사업 제안요청서를 최대한 반영할것 

                [🚨 절대 준수 지시사항 🚨]
                전체 목차를 모두 작성해서는 절대 안 됨. 
                오직 사용자가 선택한 **[{section_choice}]** 파트 하나만 집중적으로 매우 상세하고 길게 작성할 것.
                선택된 파트 이외의 다른 목차는 절대 출력하지 말것.
                표(Table) 작성 시 셀 내부에 줄바꿈 기호(\n)나 `<br>` 등의 HTML 태그를 절대 사용하지 마십시오. 여러 항목을 나열할 때는 줄바꿈 없이 쉼표(,)나 마침표(.)로만 이어 쓸것.
                또한 시스템 구축, 정보화 사업, 시스템 개발, 정보화사업 컨설팅 및 감리용역, 개인정보보호영향평가 등 ICT 사업에만 요구사항 명세를 COR-00, DAR-00 등 SW 가이드라인에 맞추어 작성할것.
                """
                final_prompt = f"{SYSTEM_PROMPT}\n\n[이번 사업 핵심 요청사항]\n{user_input}"

                result_text = generate_draft(API_KEY, final_prompt, uploaded_files)

                result_text = result_text.replace("<br>", " ").replace("<br/>", " ").replace("</br>", " ")
                
                st.markdown("### 🔍 HUG 과거 유사 사업 Top 5")
                st.caption(f"총 {len(reference_rfps_dict)}개의 과거 데이터 기반으로 현재 구상 중인 사업과 가장 유사한 레퍼런스를 추천합니다.")
                
                if reference_rfps_dict:
                    analysis_query = user_input + "\n" + result_text
                    similarity_df = get_top_5_similar_rfps_rag(analysis_query, reference_rfps_dict, API_KEY)
                    
                    if similarity_df is not None and not similarity_df.empty:
                        # 화면에는 '원본파일명' 컬럼을 숨기고 깔끔하게 표출
                        st.dataframe(similarity_df[['유사 과거 사업명', '유사도(%)']], use_container_width=True)
                        
                        # 유사도 결과 표 바로 아래에 원본 PDF 다운로드 버튼 생성
                        st.markdown("#### 📥 레퍼런스 원본 파일 다운로드")
                        for idx, row in similarity_df.iterrows():
                            file_name = row['원본파일명']
                            file_path = os.path.join("reference_rfps", file_name)
                            
                            if os.path.exists(file_path):
                                with open(file_path, "rb") as pdf_file:
                                    st.download_button(
                                        label=f"📄 {file_name} 다운로드",
                                        data=pdf_file,
                                        file_name=file_name,
                                        mime="application/pdf",
                                        key=f"download_pdf_{idx}" # 버튼 식별을 위한 고유 키
                                    )
                    else:
                        st.warning("유사도 분석에 실패했습니다.")
                else:
                    st.error("⚠️ 폴더를 찾을 수 없거나 PDF 파일이 없습니다.")
                
                st.divider() # 가로 구분
                
                # 워드 파일 생성 로직
                doc = Document()
                doc.add_heading(f"제안요청서 - {section_choice}", 0)
                doc.add_paragraph(result_text)
                
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                
                col_title, col_btn = st.columns([4, 1])
                with col_title:
                    st.markdown(f"### 📄 생성된 초안 ({section_choice})")
                with col_btn:
                    st.download_button(
                        label="💾 워드(.docx)로 다운로드", 
                        data=bio, 
                        file_name=f"제안요청서_{section_choice[:2]}.docx", 
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                        use_container_width=True
                    )
                
                st.info(result_text)
                
            except Exception as e:
                st.error(f"오류가 발생했습니다: {e}")
